"""
generate_template.py
Generiert eine vollständige Word-Vorlage (.docx) für die Masterarbeit
gemäß den Vorgaben des Leitfadens der Hochschule Osnabrück (WiSo).

Verwendung:
    pip install -r requirements.txt
    python generate_template.py

Ausgabe: Masterarbeit_Vorlage.docx
"""

import io
import os
import tempfile

import requests
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor, Inches
from docx.enum.section import WD_SECTION, WD_ORIENT

# ---------------------------------------------------------------------------
# Farben
# ---------------------------------------------------------------------------
COLOR_HS_BLUE = RGBColor(0x00, 0x75, 0xBF)
COLOR_GRAY = RGBColor(0x66, 0x66, 0x66)
COLOR_DARK_GRAY = RGBColor(0x33, 0x33, 0x33)
COLOR_BLACK = RGBColor(0x00, 0x00, 0x00)

# Mindestgröße in Bytes, um sicherzustellen, dass das Logo vollständig heruntergeladen wurde
MIN_LOGO_SIZE_BYTES = 1000

# ---------------------------------------------------------------------------
# Hilfsfunktionen – XML / OXml
# ---------------------------------------------------------------------------

def _set_para_spacing(para, before_pt=0, after_pt=0, line_rule=None, line_val=None):
    """Setzt Absatz-Abstand und Zeilenabstand direkt über pPr/pPrChange."""
    pPr = para._p.get_or_add_pPr()
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:before'), str(int(before_pt * 20)))
    spacing.set(qn('w:after'), str(int(after_pt * 20)))
    if line_rule and line_val:
        spacing.set(qn('w:lineRule'), line_rule)
        spacing.set(qn('w:line'), str(line_val))
    # Remove existing spacing element
    for existing in pPr.findall(qn('w:spacing')):
        pPr.remove(existing)
    pPr.append(spacing)


def _set_run_font(run, name, size_pt, bold=False, italic=False, color=None):
    run.font.name = name
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    # Ensure East-Asian / complex-script font also set
    rPr = run._r.get_or_add_rPr()
    for tag in ('w:rFonts',):
        existing = rPr.find(qn(tag))
        if existing is None:
            rFonts = OxmlElement('w:rFonts')
            rFonts.set(qn('w:ascii'), name)
            rFonts.set(qn('w:hAnsi'), name)
            rFonts.set(qn('w:cs'), name)
            rPr.insert(0, rFonts)
        else:
            existing.set(qn('w:ascii'), name)
            existing.set(qn('w:hAnsi'), name)
            existing.set(qn('w:cs'), name)


def _add_horizontal_rule(doc, color_rgb=None, width_pt=1):
    """Fügt eine horizontale Linie (Paragraphen-Border) ein."""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), str(width_pt * 4))  # size in 1/8 pt
    bottom.set(qn('w:space'), '1')
    if color_rgb:
        bottom.set(qn('w:color'), '{:02X}{:02X}{:02X}'.format(*color_rgb))
    else:
        bottom.set(qn('w:color'), '0075BF')
    pBdr.append(bottom)
    pPr.append(pBdr)
    _set_para_spacing(p, 0, 0)
    return p


def _set_cell_border_none(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'), 'none')
        tcBorders.append(border)
    tcPr.append(tcBorders)


def _insert_section_break(doc, break_type='nextPage'):
    """Fügt einen Abschnittsumbruch ein (nextPage oder continuous)."""
    para = doc.add_paragraph()
    pPr = para._p.get_or_add_pPr()
    sectPr = OxmlElement('w:sectPr')
    pgSz = OxmlElement('w:pgSz')
    pgSz.set(qn('w:w'), '11906')   # A4 width in twips
    pgSz.set(qn('w:h'), '16838')   # A4 height in twips
    sectPr.append(pgSz)
    pgMar = OxmlElement('w:pgMar')
    pgMar.set(qn('w:top'), '1418')    # 2.5 cm
    pgMar.set(qn('w:right'), '1418')  # 2.5 cm
    pgMar.set(qn('w:bottom'), '1418') # 2.5 cm
    pgMar.set(qn('w:left'), '1984')   # 3.5 cm
    pgMar.set(qn('w:header'), '708')
    pgMar.set(qn('w:footer'), '708')
    sectPr.append(pgMar)
    pgType = OxmlElement('w:type')
    pgType.set(qn('w:val'), break_type)
    sectPr.insert(0, pgType)
    pPr.append(sectPr)
    return para


def _set_page_number_format(sectPr, fmt='decimal', start=None):
    """Setzt das Seitenzahlformat und optionalen Startwert in einem sectPr."""
    pgNumType = OxmlElement('w:pgNumType')
    pgNumType.set(qn('w:fmt'), fmt)
    if start is not None:
        pgNumType.set(qn('w:start'), str(start))
    # remove existing
    for existing in sectPr.findall(qn('w:pgNumType')):
        sectPr.remove(existing)
    sectPr.append(pgNumType)


def _add_centered_page_number_to_footer(section):
    """Fügt eine zentrierte Seitenzahl in die Fußzeile ein."""
    footer = section.footer
    footer.is_linked_to_previous = False
    # Clear existing paragraphs
    for para in footer.paragraphs:
        p = para._p
        p.getparent().remove(p)
    para = footer.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run()
    # Insert PAGE field
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = ' PAGE '
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)


def _add_header_text(section, text, size_pt=8):
    """Fügt Text rechtsbündig in die Kopfzeile ein."""
    header = section.header
    header.is_linked_to_previous = False
    for para in header.paragraphs:
        p = para._p
        p.getparent().remove(p)
    para = header.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = para.add_run(text)
    run.font.size = Pt(size_pt)
    run.font.name = 'Times New Roman'


# ---------------------------------------------------------------------------
# Logo-Download
# ---------------------------------------------------------------------------

LOGO_URLS = [
    'https://upload.wikimedia.org/wikipedia/commons/thumb/1/1f/HS_Osnabrueck_logo.svg/1200px-HS_Osnabrueck_logo.svg.png',
    'https://www.hs-osnabrueck.de/typo3conf/ext/hs_base/Resources/Public/Images/logo.svg',
]


def _download_logo():
    """Versucht das Logo herunterzuladen. Gibt Pfad zu tmp-Datei zurück oder None."""
    headers = {'User-Agent': 'Mozilla/5.0'}
    for url in LOGO_URLS:
        try:
            resp = requests.get(url, headers=headers, timeout=10)
            if resp.status_code == 200 and len(resp.content) > MIN_LOGO_SIZE_BYTES:
                suffix = '.png' if url.endswith('.png') else '.svg'
                tmp = tempfile.NamedTemporaryFile(delete=False, suffix=suffix)
                tmp.write(resp.content)
                tmp.close()
                print(f'Logo heruntergeladen von: {url}')
                return tmp.name
        except Exception as e:
            print(f'Logo-Download fehlgeschlagen ({url}): {e}')
    return None


# ---------------------------------------------------------------------------
# Style-Definitionen
# ---------------------------------------------------------------------------

def _configure_styles(doc):
    """Konfiguriert alle benötigten Styles im Dokument."""
    styles = doc.styles

    # --- Normal ---
    normal = styles['Normal']
    normal.font.name = 'Times New Roman'
    normal.font.size = Pt(12)
    nPPr = normal.paragraph_format
    nPPr.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    nPPr.space_before = Pt(0)
    nPPr.space_after = Pt(6)
    nPPr.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

    # --- Heading 1 ---
    h1 = styles['Heading 1']
    h1.font.name = 'Times New Roman'
    h1.font.size = Pt(16)
    h1.font.bold = True
    h1.font.italic = False
    h1.font.color.rgb = COLOR_BLACK
    h1.paragraph_format.space_before = Pt(24)
    h1.paragraph_format.space_after = Pt(12)
    h1.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    # --- Heading 2 ---
    h2 = styles['Heading 2']
    h2.font.name = 'Times New Roman'
    h2.font.size = Pt(14)
    h2.font.bold = True
    h2.font.italic = False
    h2.font.color.rgb = COLOR_BLACK
    h2.paragraph_format.space_before = Pt(18)
    h2.paragraph_format.space_after = Pt(8)
    h2.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    # --- Heading 3 ---
    h3 = styles['Heading 3']
    h3.font.name = 'Times New Roman'
    h3.font.size = Pt(12)
    h3.font.bold = True
    h3.font.italic = False
    h3.font.color.rgb = COLOR_BLACK
    h3.paragraph_format.space_before = Pt(12)
    h3.paragraph_format.space_after = Pt(6)
    h3.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    # --- Footer / Footnote Text ---
    try:
        fn_style = styles['Footnote Text']
    except KeyError:
        fn_style = styles.add_style('Footnote Text', 1)
    fn_style.font.name = 'Times New Roman'
    fn_style.font.size = Pt(8)
    fn_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    fn_style.paragraph_format.space_after = Pt(0)

    # --- Blockzitat ---
    try:
        bq_style = styles['Quote']
    except KeyError:
        bq_style = styles.add_style('Quote', 1)
    bq_style.font.name = 'Times New Roman'
    bq_style.font.size = Pt(10)
    bq_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    bq_style.paragraph_format.left_indent = Cm(1)
    bq_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    bq_style.paragraph_format.space_before = Pt(6)
    bq_style.paragraph_format.space_after = Pt(6)

    # Add alias 'Blockzitat' pointing to Quote
    try:
        bz_style = styles.add_style('Blockzitat', 1)
        bz_style.base_style = bq_style
    except Exception:
        pass


# ---------------------------------------------------------------------------
# Seitenränder setzen
# ---------------------------------------------------------------------------

def _set_page_margins(section):
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.left_margin = Cm(3.5)
    section.header_distance = Cm(1.25)
    section.footer_distance = Cm(1.25)


# ---------------------------------------------------------------------------
# Deckblatt
# ---------------------------------------------------------------------------

def _add_cover_page(doc, logo_path):
    """Erstellt das Deckblatt in Section 1."""

    def centered_para(text='', size=12, bold=False, italic=False,
                      color=None, space_before=0, space_after=6):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_para_spacing(p, space_before, space_after)
        if text:
            run = p.add_run(text)
            _set_run_font(run, 'Times New Roman', size, bold=bold,
                         italic=italic, color=color)
        return p

    # Logo
    if logo_path and os.path.exists(logo_path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_para_spacing(p, 0, 6)
        try:
            run = p.add_run()
            run.add_picture(logo_path, width=Cm(6))
        except Exception:
            run.text = 'HOCHSCHULE OSNABRÜCK  ·  UNIVERSITY OF APPLIED SCIENCES'
            run.font.name = 'Times New Roman'
            run.font.size = Pt(14)
            run.font.bold = True
            run.font.color.rgb = COLOR_HS_BLUE
    else:
        # Textfallback
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_para_spacing(p, 0, 4)
        r1 = p.add_run('HOCHSCHULE OSNABRÜCK')
        _set_run_font(r1, 'Times New Roman', 18, bold=True, color=COLOR_HS_BLUE)
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_para_spacing(p2, 0, 6)
        r2 = p2.add_run('UNIVERSITY OF APPLIED SCIENCES')
        _set_run_font(r2, 'Times New Roman', 11, color=COLOR_HS_BLUE)

    # Blaue Trennlinie
    _add_horizontal_rule(doc, (0x00, 0x75, 0xBF), width_pt=2)

    # Leerzeile
    centered_para('', space_before=0, space_after=4)

    # Fakultät
    centered_para(
        'Fakultät für Wirtschafts- und Sozialwissenschaften (WiSo)',
        size=12, color=COLOR_GRAY, space_before=0, space_after=3
    )
    # Studiengang
    centered_para(
        'Studiengang Business Management (M. A.)',
        size=12, color=COLOR_GRAY, space_before=0, space_after=0
    )

    # Großer Abstand
    for _ in range(3):
        centered_para('', space_before=0, space_after=6)

    # "Masterarbeit"
    centered_para('Masterarbeit', size=20, bold=True,
                  color=COLOR_DARK_GRAY, space_before=0, space_after=10)

    # Blaue Trennlinie
    _add_horizontal_rule(doc, (0x00, 0x75, 0xBF), width_pt=1)

    # Leerzeile
    centered_para('', space_before=0, space_after=6)

    # Titel
    centered_para(
        'Multi-Agent-Systeme zur Prozessautomatisierung und Entscheidungsunterstützung'
        ' durch AI-Assisted Development',
        size=16, bold=True, space_before=0, space_after=8
    )

    # Untertitel
    centered_para(
        'Eine iterative Design-Science-Studie am Beispiel eines KI-gestützten Support-Systems',
        size=12, italic=True, color=COLOR_GRAY, space_before=0, space_after=0
    )

    # Großer Abstand
    for _ in range(5):
        centered_para('', space_before=0, space_after=6)

    # Blaue Trennlinie
    _add_horizontal_rule(doc, (0x00, 0x75, 0xBF), width_pt=1)

    # Informationsblock als Tabelle ohne Rahmen
    info = [
        ('Verfasser:', 'Nikolai Toft Loosveld'),
        ('Matrikelnummer:', '[Matrikelnummer]'),
        ('Erstprüfer:', 'Prof. Dr. Andreas Faatz'),
        ('Zweitprüfer:', 'Prof. Dr. Frank Bensberg'),
        ('Abgabedatum:', '[Datum]'),
    ]
    table = doc.add_table(rows=len(info), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = 'Table Grid'

    for i, (label, value) in enumerate(info):
        row = table.rows[i]
        # Label cell
        cell_l = row.cells[0]
        cell_l.text = ''
        run_l = cell_l.paragraphs[0].add_run(label)
        _set_run_font(run_l, 'Times New Roman', 12, bold=True)
        cell_l.paragraphs[0].paragraph_format.space_after = Pt(3)
        # Value cell
        cell_v = row.cells[1]
        cell_v.text = ''
        run_v = cell_v.paragraphs[0].add_run(value)
        _set_run_font(run_v, 'Times New Roman', 12)
        cell_v.paragraphs[0].paragraph_format.space_after = Pt(3)
        # Remove borders
        _set_cell_border_none(cell_l)
        _set_cell_border_none(cell_v)

    # Remove outer table border too
    tblPr = table._tbl.tblPr
    tblBorders = OxmlElement('w:tblBorders')
    for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'none')
        tblBorders.append(b)
    for existing in tblPr.findall(qn('w:tblBorders')):
        tblPr.remove(existing)
    tblPr.append(tblBorders)

    # Leerzeile nach Tabelle
    doc.add_paragraph()


# ---------------------------------------------------------------------------
# Seitennummerierungs-Helper
# ---------------------------------------------------------------------------

def _get_last_section_pr(doc):
    """Gibt das letzte sectPr-Element im Body zurück (Dokument-Section)."""
    body = doc.element.body
    return body.find(qn('w:sectPr'))


def _insert_section_break_with_props(doc, page_num_fmt, page_num_start=None):
    """
    Fügt einen nextPage-Abschnittsumbruch ein und setzt das Seitenzahl-Format
    für den *neuen* Abschnitt (der nach dem Bruch beginnt).
    Gibt das sectPr-Element zurück, damit der Aufrufer Footer etc. setzen kann.
    """
    # Wir fügen einen leeren Absatz mit sectPr ein
    para = doc.add_paragraph()
    pPr = para._p.get_or_add_pPr()
    sectPr = OxmlElement('w:sectPr')

    # Seitenformat A4
    pgSz = OxmlElement('w:pgSz')
    pgSz.set(qn('w:w'), '11906')
    pgSz.set(qn('w:h'), '16838')
    sectPr.append(pgSz)

    # Ränder
    pgMar = OxmlElement('w:pgMar')
    pgMar.set(qn('w:top'), '1418')
    pgMar.set(qn('w:right'), '1418')
    pgMar.set(qn('w:bottom'), '1418')
    pgMar.set(qn('w:left'), '1984')
    pgMar.set(qn('w:header'), '708')
    pgMar.set(qn('w:footer'), '708')
    sectPr.append(pgMar)

    # Typ: neue Seite
    pgType = OxmlElement('w:type')
    pgType.set(qn('w:val'), 'nextPage')
    sectPr.insert(0, pgType)

    # Seitenzahl-Format
    _set_page_number_format(sectPr, fmt=page_num_fmt, start=page_num_start)

    pPr.append(sectPr)
    return sectPr


# ---------------------------------------------------------------------------
# Abschnitt 2: Verzeichnisse
# ---------------------------------------------------------------------------

def _add_placeholder(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    _set_run_font(run, 'Times New Roman', 12, italic=True, color=COLOR_GRAY)
    _set_para_spacing(p, 0, 6,
                      line_rule='auto', line_val=360)  # 1.5-zeilig


def _add_section2_content(doc):
    """Fügt Abstract, Verzeichnisse in Section 2 ein."""

    # --- Abstract ---
    p = doc.add_heading('Abstract', level=1)
    p.runs[0].font.name = 'Times New Roman'
    p.runs[0].font.color.rgb = COLOR_BLACK

    _add_placeholder(doc, '[Zusammenfassung der Arbeit auf Deutsch, ca. 150–250 Wörter]')

    doc.add_paragraph()

    p2 = doc.add_heading('Abstract (English)', level=2)
    if p2.runs:
        p2.runs[0].font.name = 'Times New Roman'
        p2.runs[0].font.color.rgb = COLOR_BLACK

    _add_placeholder(doc, '[Summary of the thesis in English, approx. 150–250 words]')

    # Seitenumbruch
    doc.add_page_break()

    # --- Inhaltsverzeichnis ---
    p = doc.add_heading('Inhaltsverzeichnis', level=1)
    if p.runs:
        p.runs[0].font.name = 'Times New Roman'
        p.runs[0].font.color.rgb = COLOR_BLACK

    _add_placeholder(
        doc,
        '[Automatisch generiertes Inhaltsverzeichnis – '
        'In Word: Verweise → Inhaltsverzeichnis einfügen, '
        'dann Rechtsklick → Felder aktualisieren]'
    )
    doc.add_page_break()

    # --- Abkürzungsverzeichnis ---
    p = doc.add_heading('Abkürzungsverzeichnis', level=1)
    if p.runs:
        p.runs[0].font.name = 'Times New Roman'
        p.runs[0].font.color.rgb = COLOR_BLACK

    abbreviations = [
        ('AI', 'Artificial Intelligence'),
        ('DSR', 'Design Science Research'),
        ('LLM', 'Large Language Model'),
        ('MAS', 'Multi-Agent-System'),
        ('RAG', 'Retrieval Augmented Generation'),
        ('RPA', 'Robotic Process Automation'),
    ]
    table = doc.add_table(rows=len(abbreviations), cols=2)
    table.style = 'Table Grid'
    for i, (abbr, full) in enumerate(abbreviations):
        row = table.rows[i]
        cell_a = row.cells[0]
        cell_a.text = ''
        r_a = cell_a.paragraphs[0].add_run(abbr)
        _set_run_font(r_a, 'Times New Roman', 12, bold=True)
        cell_b = row.cells[1]
        cell_b.text = ''
        r_b = cell_b.paragraphs[0].add_run(full)
        _set_run_font(r_b, 'Times New Roman', 12)
        _set_cell_border_none(cell_a)
        _set_cell_border_none(cell_b)

    tblPr = table._tbl.tblPr
    tblBorders = OxmlElement('w:tblBorders')
    for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'none')
        tblBorders.append(b)
    for existing in tblPr.findall(qn('w:tblBorders')):
        tblPr.remove(existing)
    tblPr.append(tblBorders)

    doc.add_page_break()

    # --- Abbildungsverzeichnis ---
    p = doc.add_heading('Abbildungsverzeichnis', level=1)
    if p.runs:
        p.runs[0].font.name = 'Times New Roman'
        p.runs[0].font.color.rgb = COLOR_BLACK
    _add_placeholder(doc, '[Abbildungsverzeichnis – In Word automatisch generieren]')
    doc.add_page_break()

    # --- Tabellenverzeichnis ---
    p = doc.add_heading('Tabellenverzeichnis', level=1)
    if p.runs:
        p.runs[0].font.name = 'Times New Roman'
        p.runs[0].font.color.rgb = COLOR_BLACK
    _add_placeholder(doc, '[Tabellenverzeichnis – In Word automatisch generieren]')


# ---------------------------------------------------------------------------
# Abschnitt 3: Textteil
# ---------------------------------------------------------------------------

CHAPTERS = [
    ('1', 'Einleitung', 1, [
        ('1.1', 'Problemstellung und Motivation', 2),
        ('1.2', 'Zielsetzung und Forschungsfragen', 2),
        ('1.3', 'Methodisches Vorgehen und Aufbau der Arbeit', 2),
    ]),
    ('2', 'Theoretische Grundlagen', 1, [
        ('2.1', 'Geschäftsprozessautomatisierung', 2, [
            ('2.1.1', 'Begriffsdefinition und Entwicklung', 3),
            ('2.1.2', 'Herausforderungen bei wissensintensiven Prozessen', 3),
        ]),
        ('2.2', 'Multi-Agent-Systeme', 2, [
            ('2.2.1', 'Definition und Architekturprinzipien', 3),
            ('2.2.2', 'Orchestrierungsmuster und Agenten-Kommunikation', 3),
            ('2.2.3', 'Anwendungsfelder in Unternehmen', 3),
        ]),
        ('2.3', 'AI-Assisted Development', 2, [
            ('2.3.1', 'Begriffsdefinition und Abgrenzung', 3),
            ('2.3.2', 'Potenziale für Fachkräfte ohne IT-Hintergrund', 3),
        ]),
        ('2.4', 'Technologische Grundlagen', 2, [
            ('2.4.1', 'Large Language Models und RAG', 3),
            ('2.4.2', 'n8n als Workflow-Automation-Plattform', 3),
            ('2.4.3', 'Lovable als AI-gestützte Full-Stack-Entwicklungsplattform', 3),
        ]),
    ]),
    ('3', 'Forschungsmethodik', 1, [
        ('3.1', 'Design Science Research', 2, [
            ('3.1.1', 'Grundlagen und DSR-Zyklus', 3),
            ('3.1.2', 'Anwendung in dieser Arbeit', 3),
        ]),
        ('3.2', 'Anforderungsanalyse mit Volere', 2),
        ('3.3', 'Evaluationsmethodik', 2),
    ]),
    ('4', 'Anforderungsanalyse und Systemkonzeption', 1, [
        ('4.1', 'Problemkontext: Herausforderungen im Kundensupport', 2),
        ('4.2', 'Stakeholder und deren Anforderungen', 2),
        ('4.3', 'Anforderungsdefinition nach Volere', 2, [
            ('4.3.1', 'Funktionale Anforderungen', 3),
            ('4.3.2', 'Nicht-funktionale Anforderungen', 3),
        ]),
        ('4.4', 'Systemkonzeption', 2, [
            ('4.4.1', 'Gesamtarchitektur und Technologie-Stack', 3),
            ('4.4.2', 'Konzeption der Multi-Agent-Architektur als Pipeline mit dynamischem Routing', 3),
        ]),
    ]),
    ('5', 'Entwicklung des Systems', 1, [
        ('5.1', 'Überblick: Iterativer Entwicklungsprozess', 2),
        ('5.2', 'Support Hub: Frontend und Backend', 2, [
            ('5.2.1', 'Prompt-basierte Entwicklung mit Lovable', 3),
            ('5.2.2', 'Kernmodule und KI-Funktionen', 3),
        ]),
        ('5.3', 'Multi-Agent-Pipeline in n8n', 2, [
            ('5.3.1', 'Workflow-Architektur und Phasen', 3),
            ('5.3.2', 'Die fünf KI-Agenten', 3),
            ('5.3.3', 'Integration externer APIs', 3),
            ('5.3.4', 'Finaler Output: Gmail Draft und Google Task', 3),
        ]),
        ('5.4', 'Reflexion: Skills und Grenzen als Nicht-IT-Fachkraft', 2),
    ]),
    ('6', 'Evaluation', 1, [
        ('6.1', 'Vorgehen und Testszenarien', 2),
        ('6.2', 'Anforderungserfüllung (Soll-Ist-Vergleich)', 2),
        ('6.3', 'Nutzerevaluation: Funktionalität, Usability, Mehrwert', 2),
        ('6.4', 'Bewertung von AI-Assisted Development', 2),
    ]),
    ('7', 'Betriebswirtschaftliche Implikationen', 1, [
        ('7.1', 'Potenzialanalyse: Effizienz, Qualität und Skalierbarkeit', 2),
        ('7.2', 'Organisatorische Veränderungen und neue Rollen', 2),
        ('7.3', 'Erfolgsfaktoren und Handlungsempfehlungen', 2),
    ]),
    ('8', 'Schlussbetrachtung', 1, [
        ('8.1', 'Zusammenfassung und Beantwortung der Forschungsfragen', 2),
        ('8.2', 'Kritische Würdigung und Limitationen', 2),
        ('8.3', 'Ausblick', 2),
    ]),
]


def _add_chapter_recursive(doc, items):
    """Rekursiv Kapitel und Unterkapitel einfügen."""
    for item in items:
        num = item[0]
        title = item[1]
        level = item[2]
        sub = item[3] if len(item) > 3 else []

        heading_text = f'{num} {title}'
        p = doc.add_heading(heading_text, level=level)
        if p.runs:
            p.runs[0].font.name = 'Times New Roman'
            p.runs[0].font.color.rgb = COLOR_BLACK
            p.runs[0].font.bold = True

        _add_placeholder(doc, f'[Inhalt für Kapitel {num} einfügen]')

        if sub:
            _add_chapter_recursive(doc, sub)


def _add_section3_content(doc):
    """Fügt den Textteil (Kapitel 1–8) und Anhang/Erklärung ein."""

    # Kapitel 1–8
    _add_chapter_recursive(doc, CHAPTERS)

    # Seitenumbruch vor Literaturverzeichnis
    doc.add_page_break()

    # Literaturverzeichnis
    p = doc.add_heading('Literaturverzeichnis', level=1)
    if p.runs:
        p.runs[0].font.name = 'Times New Roman'
        p.runs[0].font.color.rgb = COLOR_BLACK
    _add_placeholder(
        doc,
        '[Literaturverzeichnis gemäß APA 7th Edition – '
        'alphabetisch geordnet nach Erstautor]'
    )

    # Anhang
    doc.add_page_break()
    p = doc.add_heading('Anhang', level=1)
    if p.runs:
        p.runs[0].font.name = 'Times New Roman'
        p.runs[0].font.color.rgb = COLOR_BLACK

    _add_placeholder(doc, '[Anhangsverzeichnis]')

    for label in [
        'Anhang A: [Titel]',
        'Anhang B: [Titel]',
        'Anhang C: [Titel]',
        'Anhang D: Dokumentation des KI-Einsatzes',
    ]:
        p_sub = doc.add_heading(label, level=2)
        if p_sub.runs:
            p_sub.runs[0].font.name = 'Times New Roman'
            p_sub.runs[0].font.color.rgb = COLOR_BLACK
        _add_placeholder(doc, f'[Inhalt: {label}]')

    # Eidesstattliche Erklärung
    doc.add_page_break()
    p = doc.add_heading('Eidesstattliche Erklärung', level=1)
    if p.runs:
        p.runs[0].font.name = 'Times New Roman'
        p.runs[0].font.color.rgb = COLOR_BLACK

    erklaerung_text = (
        'Hiermit erkläre ich an Eides statt, dass ich die vorliegende Arbeit '
        'selbständig und ohne fremde Hilfe angefertigt habe. Die aus fremden Quellen '
        'direkt oder indirekt übernommenen Gedanken sind als solche einzeln kenntlich '
        'gemacht. Es wurden keine anderen als die angegebenen Quellen und Hilfsmittel '
        'benutzt. Die Arbeit wurde bisher keiner anderen Prüfungsbehörde vorgelegt '
        'und auch nicht veröffentlicht.'
    )
    p_erkl = doc.add_paragraph()
    p_erkl.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run_e = p_erkl.add_run(erklaerung_text)
    _set_run_font(run_e, 'Times New Roman', 12)
    _set_para_spacing(p_erkl, 0, 6, line_rule='auto', line_val=360)

    doc.add_paragraph()

    p_ki = doc.add_paragraph()
    p_ki.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run_ki = p_ki.add_run(
        'Die eingesetzten KI-gestützten Werkzeuge sind in Anhang D dokumentiert.'
    )
    _set_run_font(run_ki, 'Times New Roman', 12)
    _set_para_spacing(p_ki, 0, 6, line_rule='auto', line_val=360)

    for _ in range(3):
        doc.add_paragraph()

    p_ort = doc.add_paragraph()
    run_ort = p_ort.add_run('Osnabrück, den [Datum]')
    _set_run_font(run_ort, 'Times New Roman', 12)
    _set_para_spacing(p_ort, 0, 6)

    doc.add_paragraph()

    p_line = doc.add_paragraph()
    run_line = p_line.add_run('_________________________________')
    _set_run_font(run_line, 'Times New Roman', 12)
    _set_para_spacing(p_line, 0, 2)

    p_name = doc.add_paragraph()
    run_name = p_name.add_run('Nikolai Toft Loosveld')
    _set_run_font(run_name, 'Times New Roman', 12)
    _set_para_spacing(p_name, 0, 0)


# ---------------------------------------------------------------------------
# Haupt-Funktion
# ---------------------------------------------------------------------------

def generate_template(output_path='Masterarbeit_Vorlage.docx'):
    print('Starte Generierung der Masterarbeit-Vorlage...')

    # Logo herunterladen
    logo_path = _download_logo()

    doc = Document()

    # Styles konfigurieren
    _configure_styles(doc)

    # ------------------------------------------------------------------ #
    # Section 1: Deckblatt (keine Seitenzahl)                            #
    # ------------------------------------------------------------------ #
    section1 = doc.sections[0]
    _set_page_margins(section1)
    # Seitenzahl-Format: leer / keine Nummer
    sectPr1 = section1._sectPr
    _set_page_number_format(sectPr1, fmt='decimal', start=1)
    # Keine Fußzeile in Section 1 (Deckblatt)
    section1.footer.is_linked_to_previous = False
    # Sicherstellen, dass die Fußzeile leer ist
    for para in section1.footer.paragraphs:
        for run in para.runs:
            run.text = ''

    _add_cover_page(doc, logo_path)

    # ------------------------------------------------------------------ #
    # Abschnittsumbruch: Ende Deckblatt → Beginn Section 2               #
    # (römische Seitenzahlen, Start bei I)                               #
    # ------------------------------------------------------------------ #
    _insert_section_break_with_props(doc, page_num_fmt='upperRoman', page_num_start=1)

    # ------------------------------------------------------------------ #
    # Section 2: Verzeichnisse (römische Seitenzahlen)                   #
    # ------------------------------------------------------------------ #
    _add_section2_content(doc)

    # ------------------------------------------------------------------ #
    # Abschnittsumbruch: Ende Verzeichnisse → Beginn Section 3           #
    # (arabische Seitenzahlen, Neustart bei 1)                           #
    # ------------------------------------------------------------------ #
    _insert_section_break_with_props(doc, page_num_fmt='decimal', page_num_start=1)

    # ------------------------------------------------------------------ #
    # Section 3: Textteil + Literatur + Anhang + Erklärung               #
    # ------------------------------------------------------------------ #
    _add_section3_content(doc)

    # ------------------------------------------------------------------ #
    # Abschluss-Section (letzte sectPr im body):                         #
    # Seitenzahl arabisch fortlaufend, Kopfzeile, Fußzeile               #
    # ------------------------------------------------------------------ #
    last_sectPr = doc.element.body.find(qn('w:sectPr'))
    if last_sectPr is not None:
        _set_page_number_format(last_sectPr, fmt='decimal')
        _set_page_margins(doc.sections[-1])

    # Fußzeilen für Section 2 und 3 über python-docx sections setzen
    # sections[0] = Deckblatt (keine Nummer)
    # sections[1] = Verzeichnisse (römisch) – entsteht durch unsere sectPr-Einschübe
    # sections[2] = Textteil (arabisch)
    # Hinweis: python-docx's sections-Liste bildet die im Body eingebetteten sectPr ab.
    for i, section in enumerate(doc.sections):
        _set_page_margins(section)
        if i == 0:
            # Deckblatt: keine Seitenzahl
            footer = section.footer
            footer.is_linked_to_previous = False
            for para in footer.paragraphs:
                for run in para.runs:
                    run.text = ''
        else:
            # Alle anderen Sections: Seitenzahl in Fußzeile
            _add_centered_page_number_to_footer(section)
        if i >= 2:
            # Textteil: Kopfzeile mit Name
            _add_header_text(
                section, 'Masterarbeit – Nikolai Toft Loosveld', size_pt=8
            )

    # Datei speichern
    doc.save(output_path)
    print(f'Vorlage erfolgreich gespeichert: {output_path}')

    # Temp-Logo aufräumen
    if logo_path and os.path.exists(logo_path):
        try:
            os.unlink(logo_path)
        except Exception:
            pass


if __name__ == '__main__':
    generate_template()
